from docx import Document
import os
from datetime import datetime

class DocGenerationService:
    def __init__(self):
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)

    async def create_analysis_document(self, analysis: str) -> str:
        """Create a Word document with the analysis."""
        doc = Document()
        
        # Add title
        doc.add_heading('Competitor Analysis Report', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add content
        sections = analysis.split('\n\n')
        for section in sections:
            if section.strip():
                if section.startswith('#'):
                    # Add as heading
                    doc.add_heading(section.strip('# '), level=1)
                else:
                    doc.add_paragraph(section)

        # Save document
        filename = f"competitor_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath 