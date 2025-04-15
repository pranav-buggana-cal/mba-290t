import logging
from typing import List, Dict, Tuple, Optional, Any, Union
import os
import io
import re
from concurrent.futures import ThreadPoolExecutor
import uuid
import asyncio
from openai import OpenAI
import json
import tempfile
import random

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument

# File processing imports
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Define small file threshold (in characters)
SMALL_FILE_THRESHOLD = 5000  # Skip chunking for files below 5000 characters
# Define optimal batch size for processing
OPTIMAL_BATCH_SIZE = 20  # Process 20 chunks at a time for better throughput
# Maximum batch size (hard limit to avoid overwhelming API)
MAX_BATCH_SIZE = 50


class RAGService:
    def __init__(
        self,
        openai_service,
        vector_db,
        prompt_path: str = None,
        chunk_size: int = 5000,
        chunk_overlap: int = 600,
        max_workers: int = 5,
    ):
        """Initialize the RAG Service with LangChain components for text splitting.

        Args:
            openai_service: Service for OpenAI API calls
            vector_db: Vector database service
            prompt_path: Path to prompts file
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
            max_workers: Maximum number of parallel workers for processing tasks
        """
        self.openai_service = openai_service
        self.vector_db = vector_db
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_workers = max_workers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)

        # Initialize LangChain text splitter only
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

        # Cache the original settings for later reference
        self._original_chunk_size = self.chunk_size
        self._original_chunk_overlap = self.chunk_overlap

        # Load prompts
        self.prompts = self._load_prompts(prompt_path)

    def _load_prompts(self, prompt_path: str = None):
        """Load prompt templates from file."""
        # Get the absolute path to the prompts file
        current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        prompts_path = prompt_path or os.path.join(
            current_dir, "prompts", "template_prompts.txt"
        )

        try:
            with open(prompts_path, "r") as f:
                content = f.read()
                logger.info(
                    f"Loaded prompt template file with {len(content)} characters"
                )

                self.prompts = {}

                # Split by template headers (# followed by a name)
                # This regex matches # at the start of a line followed by text
                sections = re.split(r"(?=^# .*$)", content, flags=re.MULTILINE)
                # Remove empty sections at the beginning
                sections = [s for s in sections if s.strip()]

                logger.info(f"Found {len(sections)} template sections")

                for section in sections:
                    # Extract the section name from the first line
                    lines = section.strip().split("\n")
                    if not lines:
                        continue

                    first_line = lines[0].strip()
                    if not first_line.startswith("# "):
                        logger.warning(f"Invalid header: {first_line[:30]}...")
                        continue

                    name = first_line[2:].strip()  # Remove '# ' prefix
                    # The template is everything after the first line
                    template = "\n".join(lines[1:]).strip()

                    # Log the template details
                    logger.info(
                        f"Loaded template: {name} with {len(template)} characters"
                    )
                    logger.info(f"Template preview: {template[:200]}...")

                    # For competitor analysis template, verify critical sections exist
                    if name == "Competitor Analysis Template":
                        # Ensure required markers are present
                        required_markers = [
                            "{query}",
                            "Your analysis should include",
                            "Executive Summary",
                            "Industry Analysis",
                        ]

                        for marker in required_markers:
                            if marker not in template:
                                logger.error(
                                    f"Template missing critical marker: {marker}"
                                )
                                # Fix: Reload the template from the original file
                                logger.info("Attempting to fix template...")
                                with open(prompts_path, "r") as fix_file:
                                    fix_content = fix_file.read()
                                    if marker in fix_content:
                                        logger.info(
                                            f"Marker '{marker}' found in original file"
                                        )
                                        # Extract the complete template directly
                                        match = re.search(
                                            r"^# Competitor Analysis Template$(.*?)(?=^# |\Z)",
                                            fix_content,
                                            re.MULTILINE | re.DOTALL,
                                        )
                                        if match:
                                            template = match.group(1).strip()
                                            logger.info(
                                                "Template fixed from original file"
                                            )
                                        else:
                                            logger.error(
                                                "Failed to extract template from original file"
                                            )

                    self.prompts[name] = template

                logger.info(f"Successfully loaded {len(self.prompts)} templates")

        except FileNotFoundError:
            logger.error(f"Could not find prompts file at {prompts_path}")
            raise FileNotFoundError(f"Could not find prompts file at {prompts_path}")
        except Exception as e:
            logger.error(f"Error loading prompts: {str(e)}")
            raise Exception(f"Error loading prompts: {str(e)}")

    def _extract_text_from_file(
        self, content: bytes, filename: str, content_type: str = None
    ) -> str:
        """Extract text from different file types."""
        file_extension = os.path.splitext(filename.lower())[1]

        try:
            # Try to determine file type from extension or provided content type
            if not file_extension and content_type:
                # Map content types to extensions
                content_type_map = {
                    "application/pdf": ".pdf",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
                    "text/plain": ".txt",
                }
                file_extension = content_type_map.get(content_type, "")
                logger.info(
                    f"Determined file extension from content type: {file_extension}"
                )

            if file_extension == ".pdf" or content_type == "application/pdf":
                # Handle PDF files
                pdf_file = io.BytesIO(content)
                pdf_reader = PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text

            elif (
                file_extension == ".docx"
                or content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                # Handle DOCX files
                doc_file = io.BytesIO(content)
                doc = DocxDocument(doc_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text

            elif file_extension == ".txt" or content_type == "text/plain":
                # Handle plain text files
                return content.decode("utf-8")

            else:
                raise ValueError(
                    f"Unsupported file type: "
                    f"{file_extension or content_type or 'unknown'}"
                )

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        file_type: str = None,
        content_type: str = None,
    ):
        """Process a document using LangChain for chunking.

        Args:
            file_content: Binary content of the file
            filename: Name of the file
            file_type: Type of the file ('competitor' or 'business')
            content_type: MIME type of the file content

        Returns:
            Dictionary with processing results
        """
        try:
            logger.info(f"Processing document: {filename}, type: {file_type}")
            file_size = len(file_content)
            logger.info(f"File size: {file_size} bytes")

            # Extract text from document
            text = self._extract_text_from_file(file_content, filename, content_type)
            if not text:
                logger.error(f"Failed to extract text from {filename}")
                raise ValueError(f"No text content extracted from {filename}")

            # Calculate text stats for logging
            text_length = len(text)
            line_count = text.count("\n") + 1
            logger.info(
                f"Extracted {text_length} characters, {line_count} lines from {filename}"
            )

            # Handle different file sizes
            active_splitter = self.text_splitter
            if text_length > 100000:  # For extremely large files (>100K chars)
                logger.info(
                    f"Large file detected ({text_length} chars). Using smaller chunks."
                )
                # Create a temporary text splitter with smaller chunks for large files
                new_chunk_size = min(2000, self._original_chunk_size)
                new_chunk_overlap = min(100, self._original_chunk_overlap)

                active_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=new_chunk_size,
                    chunk_overlap=new_chunk_overlap,
                    length_function=len,
                    is_separator_regex=False,
                )
                logger.info(
                    f"Created temporary text splitter with smaller chunk size: {new_chunk_size} "
                    f"(original: {self._original_chunk_size})"
                )

            # Check if this is a small file that can skip chunking
            if text_length <= SMALL_FILE_THRESHOLD:
                logger.info(
                    f"Small file detected ({text_length} chars). Using fast path without chunking."
                )

                # Process the entire document as a single chunk
                metadata = {"source": filename, "doc_type": file_type or "unknown"}

                # Get embedding for the entire document
                embedding = await self.openai_service.get_embedding(text)

                # Store in vector DB
                await self.vector_db.add_document(
                    text=text,
                    embedding=embedding,
                    metadata=metadata,
                )

                logger.info(
                    f"Successfully processed small document as single chunk: {filename}"
                )

                # Return processing details
                return {
                    "status": "success",
                    "message": "Processed small document as single chunk",
                    "total_chunks": 1,
                    "processed_chunks": 1,
                }

            # For larger documents, create chunks with LangChain
            langchain_doc = LangchainDocument(
                page_content=text,
                metadata={"source": filename, "doc_type": file_type or "unknown"},
            )

            # Use the selected text splitter to create chunks
            documents = active_splitter.split_documents([langchain_doc])
            total_chunks = len(documents)
            logger.info(f"Created {total_chunks} chunks with LangChain")

            # Track progress
            processed_chunks = 0

            # Process chunks in batches to avoid overwhelming the API
            # For large files, use smaller batch sizes
            if text_length > 50000:  # >50K chars
                adaptive_batch_size = min(5, self.max_workers)  # Smaller batches
            else:
                # Determine optimal batch size based on document size
                adaptive_batch_size = min(
                    max(OPTIMAL_BATCH_SIZE, self.max_workers * 2),  # Scale with workers
                    MAX_BATCH_SIZE,  # But don't exceed max
                    max(1, total_chunks // 4),  # Divide into 4+ batches if possible
                )

            logger.info(
                f"Processing {total_chunks} chunks "
                f"in batches of {adaptive_batch_size}"
            )

            # Track processing progress and handle errors for large documents
            try:
                # Create asyncio tasks for batch processing
                tasks = []

                for i in range(0, total_chunks, adaptive_batch_size):
                    batch = documents[i : i + adaptive_batch_size]
                    # Use asyncio coroutine directly instead of ThreadPoolExecutor
                    tasks.append(self._process_chunk_batch(batch, file_type))

                # Wait for all tasks to complete with proper asyncio handling
                results = await asyncio.gather(*tasks)

                # Update progress
                processed_chunks = sum(results)
                progress_percent = int(processed_chunks / total_chunks * 100)
                logger.info(
                    f"Processing progress: {processed_chunks}/{total_chunks} "
                    f"chunks ({progress_percent}%)"
                )

            except Exception as e:
                logger.error(f"Error processing document: {str(e)}")
                raise

            logger.info(f"Successfully processed document: {filename}")

            # Return processing details
            return {
                "status": "success",
                "message": (
                    f"Processed {processed_chunks}/{total_chunks} chunks with LangChain"
                ),
                "total_chunks": total_chunks,
                "processed_chunks": processed_chunks,
            }

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise

    async def get_relevant_context(self, query: str) -> Tuple[List[Dict], List[Dict]]:
        """Get relevant context for a query.

        Args:
            query: The query to retrieve context for

        Returns:
            Tuple containing two lists of dictionaries:
            - Competitor documents with text and metadata
            - Business documents with text and metadata
        """
        try:
            # Get embedding for the query using our existing service
            query_embedding = await self.openai_service.get_embedding(query)

            # Retrieve relevant documents from vector DB
            results = await self.vector_db.search(query_embedding, limit=10)

            # Separate competitor and business documents
            competitor_docs = []
            business_docs = []

            for doc in results:
                doc_type = doc.get("metadata", {}).get("doc_type", "unknown")

                if doc_type == "competitor":
                    competitor_docs.append(doc)
                elif doc_type == "business":
                    business_docs.append(doc)

            logger.info(
                f"Retrieved {len(competitor_docs)} competitor and "
                f"{len(business_docs)} business documents"
            )

            return competitor_docs, business_docs

        except Exception as e:
            logger.error(f"Error retrieving context: {str(e)}")
            raise

    async def clear_vector_db(self):
        """Clear all documents from the vector database."""
        try:
            logger.info("Clearing vector database...")
            await self.vector_db.clear_collection()
            logger.info("Vector database cleared successfully")
        except Exception as e:
            logger.error(f"Error clearing vector database: {str(e)}")
            raise

    async def generate_analysis(
        self, query: str, context: Tuple[List[Dict], List[Dict]]
    ) -> str:
        """Generate analysis using RAG."""
        try:
            competitor_docs, business_docs = context

            # Join texts separately
            competitor_text = "\n\n".join([doc["text"] for doc in competitor_docs])
            business_text = "\n\n".join([doc["text"] for doc in business_docs])

            total_length = len(competitor_text) + len(business_text)
            logger.info(f"Generating analysis with context length: {total_length}")
            logger.info(f"Query: {query}")
            logger.info(f"Competitor text length: {len(competitor_text)}")
            logger.info(f"Business text length: {len(business_text)}")

            # Get the template directly from the file to avoid any potential truncation
            current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            prompts_path = os.path.join(current_dir, "prompts", "template_prompts.txt")

            try:
                with open(prompts_path, "r") as f:
                    content = f.read()
                    # Extract the competitor analysis template directly from the file
                    match = re.search(
                        r"^# Competitor Analysis Template$(.*?)(?=^# |\Z)",
                        content,
                        re.MULTILINE | re.DOTALL,
                    )
                    if match:
                        fresh_template = match.group(1).strip()
                        logger.info(
                            f"Loaded fresh template with {len(fresh_template)} chars"
                        )
                        template = fresh_template
                    else:
                        # Fall back to the cached template
                        template_key = "Competitor Analysis Template"
                        template = self.prompts[template_key]
                        logger.warning("Using cached template as fallback")
            except Exception as e:
                logger.warning(
                    f"Error loading fresh template: {e}, using cached version"
                )
                template_key = "Competitor Analysis Template"
                template = self.prompts[template_key]

            logger.info(f"Template length: {len(template)} chars")
            logger.info(f"Template preview: {template[:300]}...")
            logger.info(f"Template middle part: {template[300:600]}...")
            logger.info(f"Template end part: {template[-300:]}...")

            # Check template for required placeholders
            required_placeholders = ["{query}", "{context}", "{business_context}"]
            missing_placeholders = [
                p for p in required_placeholders if p not in template
            ]
            if missing_placeholders:
                logger.warning(f"Template missing placeholders: {missing_placeholders}")
                # Add missing placeholders to template if needed
                if "{context}" not in template:
                    logger.info("Adding {context} placeholder to template")
                    template = template.replace(
                        "Competitor Context:", "Competitor Context:\n{context}"
                    )
                if "{business_context}" not in template:
                    logger.info("Adding {business_context} placeholder to template")
                    template += "\n\nBusiness Context:\n{business_context}"

            # Format the template with the provided values
            try:
                # Final verification of template integrity
                if "Your analysis should include all of the following" in template:
                    start_idx = template.find(
                        "Your analysis should include all of the following"
                    )
                    section_requirements = template[start_idx : start_idx + 800]
                    logger.info(f"Section requirements: {section_requirements}")

                # Format the template with both competitor and business contexts
                prompt = template.format(
                    query=query, context=competitor_text, business_context=business_text
                )
                logger.info("Template formatting successful")
                logger.info(f"Final prompt length: {len(prompt)} chars")
                logger.info(f"Final prompt preview: {prompt[:200]}...")

            except Exception as e:
                logger.error(f"Template formatting failed: {e}")
                # Fall back to basic formatting
                prompt = (
                    f"You are a strategic business analyst. "
                    f"Based on the provided context about competitors "
                    f"and the specific query, provide a detailed competitive analysis "
                    f"and strategic recommendations.\n\n"
                    f"User's query: {query}\n\n"
                    f"Competitor context:\n{competitor_text}\n\n"
                    f"Business context:\n{business_text}\n\n"
                    f"Generate a comprehensive analysis with these sections:\n"
                    f"1. Executive Summary\n"
                    f"2. List of Top Competitors\n"
                    f"3. Industry Analysis\n"
                    f"4. Market Positioning\n"
                    f"5. Competitive Analysis\n"
                    f"6. Strategic Recommendations\n"
                    f"7. Risk Assessment\n"
                )
                logger.info("Used simplified fallback template")

            # Generate the analysis
            analysis = await self.openai_service.generate_completion(prompt)
            logger.info(f"Generated analysis of length: {len(analysis)}")

            return analysis
        except Exception as e:
            logger.error(f"Error generating analysis: {str(e)}")
            raise

    async def _process_chunk_batch(self, batch, file_type):
        """Process a batch of document chunks by adding them to the vector database.

        Args:
            batch: List of LangchainDocument objects to process
            file_type: Type of the document ('competitor' or 'business')

        Returns:
            Number of successfully processed chunks
        """
        try:
            logger.info(f"Processing batch of {len(batch)} chunks")

            # Extract text from each document in the batch
            texts = [doc.page_content for doc in batch]

            # Get embeddings for all texts in the batch
            embeddings = await self.openai_service.get_embeddings_batch(texts)

            # Process each document with its embedding
            processed_count = 0
            for i, (doc, embedding) in enumerate(zip(batch, embeddings)):
                try:
                    # Prepare metadata from the document
                    metadata = doc.metadata.copy()
                    if file_type and "doc_type" not in metadata:
                        metadata["doc_type"] = file_type

                    # Add to vector database
                    await self.vector_db.add_document(
                        text=doc.page_content, embedding=embedding, metadata=metadata
                    )
                    processed_count += 1
                except Exception as e:
                    logger.error(f"Error processing document chunk {i}: {str(e)}")

            logger.info(f"Successfully processed {processed_count}/{len(batch)} chunks")
            return processed_count

        except Exception as e:
            logger.error(f"Error in batch processing: {str(e)}")
            # Return 0 for fully failed batches
            return 0
